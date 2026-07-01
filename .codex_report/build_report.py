from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = ROOT / ".codex_report"
SCREEN_DIR = REPORT_DIR / "screenshots"
TEMPLATE = REPORT_DIR / "template.docx"
OUTPUT = ROOT / "BaoCao_DoAn_CMS_PhucLong_NguyenTrucTruong.docx"
ERD_PATH = REPORT_DIR / "erd-cms.png"
SWAGGER_REPORT = SCREEN_DIR / "01-swagger-ui-report.png"


def clear_document(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(11)


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "808080")


def format_run(run, size: int | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def setup_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(13)

    for name, size in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 13)]:
        style = styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)


def add_page_number(section) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)


def para(doc: Document, text: str = "", align=None, bold: bool = False, size: int | None = None, italic: bool = False):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.8) if text else None
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(6)
    if align is not None:
        p.alignment = align
    elif text:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    format_run(r, size=size, bold=bold)
    r.italic = italic
    return p


def heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    for run in p.runs:
        format_run(run, bold=True)
    return p


def bullet(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.first_line_indent = Cm(-0.4)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("• " + text)
    format_run(r, size=13)
    return p


def caption(doc: Document, text: str):
    p = para(doc, text, align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=11)
    p.paragraph_format.first_line_indent = None
    return p


def add_image(doc: Document, image_name: str | Path, cap: str, width: float = 6.2):
    path = Path(image_name)
    if not path.is_absolute():
        path = SCREEN_DIR / path
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    caption(doc, cap)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_borders(table)
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_shading(hdr_cells[i], "D9EAF7")
        set_cell_text(hdr_cells[i], header, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    if widths:
        for row in table.rows:
            for cell, width in zip(row.cells, widths):
                cell.width = Inches(width)
    doc.add_paragraph()
    return table


def draw_erd() -> None:
    W, H = 1800, 1120
    img = Image.new("RGB", (W, H), "#ffffff")
    d = ImageDraw.Draw(img)
    try:
        font_title = ImageFont.truetype("arial.ttf", 28)
        font_head = ImageFont.truetype("arialbd.ttf", 22)
        font_text = ImageFont.truetype("arial.ttf", 18)
        font_rel = ImageFont.truetype("arial.ttf", 17)
    except OSError:
        font_title = font_head = font_text = font_rel = ImageFont.load_default()

    d.text((W // 2, 28), "ERD hệ thống CMS và bán hàng Phúc Long", anchor="mm", fill="#0b6b3a", font=font_title)

    boxes = {
        "PostCategory": (70, 110, 380, 310, ["PK Id", "Name", "Slug", "ParentId"]),
        "Post": (520, 110, 860, 350, ["PK Id", "Title", "Content", "ImageUrl", "PostCategoryId", "Slug"]),
        "User": (1210, 110, 1530, 300, ["PK Id", "Username", "PasswordHash", "FullName", "Role"]),
        "ProductCategory": (70, 470, 410, 700, ["PK Id", "Name", "Slug", "ParentId", "ImageUrl"]),
        "Product": (520, 450, 880, 750, ["PK Id", "Name", "Price", "StockQuantity", "TotalSold", "ProductCategoryId", "Slug"]),
        "Customer": (70, 820, 410, 1040, ["PK Id", "FullName", "Email", "Phone", "Password", "TokenVersion"]),
        "Order": (520, 820, 910, 1050, ["PK Id", "CustomerId", "OrderDate", "Status", "ShippingFee", "DiscountAmount", "TotalAmount"]),
        "OrderDetail": (1090, 780, 1460, 1040, ["PK Id", "OrderId", "ProductId", "Quantity", "UnitPrice"]),
    }

    def draw_box(name, spec):
        x1, y1, x2, y2, fields = spec
        d.rounded_rectangle((x1, y1, x2, y2), radius=16, fill="#f8fbfa", outline="#0b6b3a", width=3)
        d.rectangle((x1, y1, x2, y1 + 44), fill="#0b6b3a")
        d.text(((x1 + x2) // 2, y1 + 23), name, anchor="mm", fill="white", font=font_head)
        y = y1 + 58
        for field in fields:
            d.text((x1 + 18, y), field, fill="#222222", font=font_text)
            y += 28

    for name, spec in boxes.items():
        draw_box(name, spec)

    def center(name, side):
        x1, y1, x2, y2, _ = boxes[name]
        if side == "r":
            return x2, (y1 + y2) // 2
        if side == "l":
            return x1, (y1 + y2) // 2
        if side == "b":
            return (x1 + x2) // 2, y2
        return (x1 + x2) // 2, y1

    relations = [
        ("PostCategory", "r", "Post", "l", "1", "N", "phân loại bài viết"),
        ("ProductCategory", "r", "Product", "l", "1", "N", "phân loại sản phẩm"),
        ("Customer", "r", "Order", "l", "1", "N", "đặt hàng"),
        ("Order", "r", "OrderDetail", "l", "1", "N", "chi tiết đơn"),
        ("Product", "b", "OrderDetail", "t", "1", "N", "được mua"),
    ]
    for a, side_a, b, side_b, la, lb, label in relations:
        x1, y1 = center(a, side_a)
        x2, y2 = center(b, side_b)
        d.line((x1, y1, x2, y2), fill="#303030", width=3)
        d.text((x1 + 8, y1 - 26), la, fill="#0b6b3a", font=font_head)
        d.text((x2 - 22, y2 - 26), lb, fill="#0b6b3a", font=font_head)
        d.text(((x1 + x2) // 2, (y1 + y2) // 2 - 22), label, fill="#303030", font=font_rel)

    d.arc((150, 250, 340, 420), start=90, end=350, fill="#707070", width=2)
    d.text((230, 380), "cây cha-con", fill="#707070", font=font_rel)
    d.arc((170, 650, 370, 820), start=90, end=350, fill="#707070", width=2)
    d.text((230, 760), "cây cha-con", fill="#707070", font=font_rel)

    img.save(ERD_PATH)


def prepare_images() -> None:
    src = SCREEN_DIR / "01-swagger-ui.png"
    if src.exists():
        im = Image.open(src)
        crop = im.crop((0, 0, im.width, min(im.height, 1850)))
        crop.save(SWAGGER_REPORT)
    draw_erd()


def build_document() -> None:
    prepare_images()
    doc = Document(str(TEMPLATE))
    clear_document(doc)
    setup_styles(doc)
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)
    add_page_number(section)

    # Cover
    for text, size in [
        ("TRƯỜNG CAO ĐẲNG CÔNG THƯƠNG TP. HỒ CHÍ MINH", 14),
        ("KHOA CÔNG NGHỆ THÔNG TIN", 14),
    ]:
        p = para(doc, text, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=size)
        p.paragraph_format.first_line_indent = None
    doc.add_paragraph()
    p = para(doc, "ĐỒ ÁN MÔN HỌC", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=20)
    p.paragraph_format.first_line_indent = None
    p = para(doc, "CHUYÊN ĐỀ ASP.NET", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=18)
    p.paragraph_format.first_line_indent = None
    doc.add_paragraph()
    p = para(
        doc,
        "XÂY DỰNG HỆ THỐNG CMS VÀ WEBSITE BÁN HÀNG PHÚC LONG\nBẰNG ASP.NET CORE WEB API VÀ REACTJS",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        bold=True,
        size=16,
    )
    p.paragraph_format.first_line_indent = None
    doc.add_paragraph()
    for text in [
        "Giảng viên hướng dẫn: ThS. NGUYỄN CAO THÁI",
        "Sinh viên thực hiện: NGUYỄN TRÚC TRƯỜNG",
        "Mã số sinh viên: 23082005",
        "Lớp: 24DTH26",
    ]:
        p = para(doc, text, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
        p.paragraph_format.first_line_indent = None
    doc.add_paragraph()
    p = para(doc, "TP. Hồ Chí Minh, Tháng 06 - Năm 2026", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=13)
    p.paragraph_format.first_line_indent = None
    doc.add_page_break()

    heading(doc, "MỤC LỤC", 1)
    toc_items = [
        "CHƯƠNG 1: TỔNG QUAN VỀ ĐỀ TÀI VÀ CƠ SỞ LÝ THUYẾT",
        "CHƯƠNG 2: PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG",
        "CHƯƠNG 3: TRIỂN KHAI BACKEND ASP.NET CORE",
        "CHƯƠNG 4: TRIỂN KHAI FRONTEND REACTJS",
        "CHƯƠNG 5: KIỂM THỬ, MINH CHỨNG VÀ ĐÁNH GIÁ",
        "CHƯƠNG 6: KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN",
        "PHỤ LỤC: ĐỐI CHIẾU TIÊU CHÍ CHẤM ĐIỂM",
    ]
    for item in toc_items:
        para(doc, item + " ........................................................................", size=13)
    doc.add_page_break()

    heading(doc, "DANH MỤC HÌNH ẢNH/SƠ ĐỒ", 1)
    figures = [
        "Hình 2.1. Sơ đồ ERD hệ thống CMS và bán hàng",
        "Hình 3.1. Swagger Web API",
        "Hình 3.2. Kết quả JSON của API sản phẩm",
        "Hình 3.3. Màn hình đăng nhập Admin",
        "Hình 3.4. Dashboard quản trị",
        "Hình 3.5. CRUD sản phẩm",
        "Hình 3.6. CRUD đơn hàng",
        "Hình 3.7. CKEditor và upload ảnh bài viết",
        "Hình 4.1. Trang chủ ReactJS",
        "Hình 4.2. Tìm kiếm/lọc sản phẩm",
        "Hình 4.3. Chi tiết sản phẩm",
        "Hình 4.4. Thêm giỏ hàng và badge số lượng",
        "Hình 4.5. Checkout",
        "Hình 4.6. Store locator và bản đồ",
        "Hình 4.7. Chi tiết bài viết HTML",
    ]
    for item in figures:
        para(doc, item, size=12)

    heading(doc, "DANH MỤC BẢNG BIỂU", 1)
    tables = [
        "Bảng 2.1. Nhóm chức năng chính của hệ thống",
        "Bảng 2.2. Thiết kế các bảng dữ liệu chính",
        "Bảng 2.3. Danh mục Web API tiêu biểu",
        "Bảng 5.1. Kết quả build và kiểm thử",
        "Bảng P.1. Đối chiếu tiêu chí chấm điểm",
    ]
    for item in tables:
        para(doc, item, size=12)
    doc.add_page_break()

    # Chapter 1
    heading(doc, "CHƯƠNG 1: TỔNG QUAN VỀ ĐỀ TÀI VÀ CƠ SỞ LÝ THUYẾT", 1)
    heading(doc, "1.1. Giới thiệu đề tài", 2)
    para(
        doc,
        "Đề tài xây dựng một hệ thống CMS kết hợp website bán hàng cho thương hiệu Phúc Long. Hệ thống gồm ba lớp rõ ràng: CMS.Data quản lý thực thể và DbContext, CMS.Backend cung cấp trang quản trị MVC/Razor cùng Web API, và CMS.Frontend là ứng dụng ReactJS phục vụ khách hàng. Cách tổ chức này giúp phần dữ liệu, nghiệp vụ và giao diện tách biệt, dễ mở rộng và dễ kiểm thử.",
    )
    para(
        doc,
        "Website mô phỏng luồng thương mại điện tử ngành đồ uống: xem banner, duyệt danh mục, tìm kiếm/lọc sản phẩm, xem chi tiết, chọn tùy chọn, thêm giỏ hàng, thanh toán, quản lý đơn hàng và đọc bài viết giới thiệu/khuyến mãi. Trang quản trị cho phép quản lý sản phẩm, danh mục, bài viết, người dùng, khách hàng, đơn hàng, cửa hàng, đánh giá và voucher.",
    )
    heading(doc, "1.2. Mục tiêu", 2)
    for item in [
        "Khởi tạo solution đúng mô hình 3 tầng và kết nối SQL Server bằng Entity Framework Core.",
        "Xây dựng backend ASP.NET Core vừa phục vụ MVC admin vừa phục vụ Web API JSON cho ReactJS.",
        "Xây dựng frontend ReactJS có giao diện bán hàng, giỏ hàng, checkout và tra cứu cửa hàng.",
        "Tích hợp xác thực, phân quyền, CORS, Swagger, upload ảnh CKEditor, email và kiểm soát tồn kho.",
        "Kiểm thử bằng build, unit test, lint/build frontend và chạy thực tế để chụp màn hình minh chứng.",
    ]:
        bullet(doc, item)
    heading(doc, "1.3. Phạm vi và đối tượng sử dụng", 2)
    para(doc, "Đối tượng sử dụng gồm quản trị viên hệ thống, nhân viên quản lý nội dung/sản phẩm và khách hàng mua hàng trên website. Phạm vi tập trung vào hệ thống quản trị CMS, API thương mại điện tử và giao diện ReactJS chạy trên môi trường localhost.")
    heading(doc, "1.4. Công nghệ sử dụng", 2)
    add_table(
        doc,
        ["Nhóm", "Công nghệ", "Vai trò"],
        [
            ["Backend", "ASP.NET Core 8, MVC, Web API", "Xử lý nghiệp vụ, trang quản trị, API JSON"],
            ["Data", "Entity Framework Core, SQL Server", "Thiết kế thực thể, migration, truy vấn dữ liệu"],
            ["Frontend", "ReactJS 18, Vite, Axios, CSS Modules", "Giao diện khách hàng, gọi API, quản lý trạng thái"],
            ["Bảo mật", "Cookie Auth, JWT Bearer, Authorize, PasswordHasher", "Đăng nhập admin/customer, phân quyền, hash mật khẩu"],
            ["Tài liệu/kiểm thử", "Swagger, xUnit, Vite build/lint", "Tài liệu API, kiểm thử tự động, xác nhận frontend"],
        ],
    )

    # Chapter 2
    heading(doc, "CHƯƠNG 2: PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG", 1)
    heading(doc, "2.1. Yêu cầu chức năng", 2)
    add_table(
        doc,
        ["Phân hệ", "Chức năng chính"],
        [
            ["Quản trị", "Đăng nhập admin, hiển thị FullName/Role, sidebar điều hướng, CRUD danh mục, bài viết, sản phẩm, khách hàng, đơn hàng, cửa hàng, voucher và người dùng."],
            ["Nội dung", "Quản lý danh mục bài viết, bài viết có slug, ảnh đại diện, nội dung HTML từ CKEditor và upload ảnh trực tiếp trong trình soạn thảo."],
            ["Sản phẩm", "Quản lý danh mục sản phẩm dạng cây, sản phẩm, ảnh, tồn kho, tổng đã bán, tùy chọn size/đường/đá và danh sách sản phẩm bán chạy/mới nhất."],
            ["Khách hàng", "Đăng ký, đăng nhập JWT, quên mật khẩu, cập nhật hồ sơ, địa chỉ, giỏ hàng và lịch sử đơn hàng."],
            ["Đơn hàng", "Tạo đơn từ frontend, kiểm tra tồn kho, tính tổng tiền, phí giao hàng, voucher, gửi email xác nhận và quản lý trạng thái đơn."],
            ["Frontend", "Trang chủ, menu/shop, tìm kiếm, lọc giá, chi tiết sản phẩm, giỏ hàng, checkout, bài viết, store locator và bản đồ."],
        ],
    )
    heading(doc, "2.2. Sơ đồ ERD", 2)
    para(doc, "Sơ đồ dưới đây thể hiện tám thực thể trọng tâm theo tiêu chí chấm điểm: PostCategory, Post, ProductCategory, Product, Customer, Order, OrderDetail và User. Dự án còn mở rộng thêm Store, Voucher, Review, Banner, ProductImage và các bảng tùy chọn sản phẩm.")
    add_image(doc, ERD_PATH, "Hình 2.1. Sơ đồ ERD hệ thống CMS và bán hàng", width=6.4)
    heading(doc, "2.3. Thiết kế cơ sở dữ liệu", 2)
    add_table(
        doc,
        ["Bảng/Entity", "Trường chính", "Quan hệ"],
        [
            ["PostCategory", "Id, Name, Description, Slug, ParentId", "1-n Post; tự liên kết cha-con"],
            ["Post", "Id, Title, Content, ImageUrl, ExternalId, SourceUrl, PublishedAt, PostCategoryId, Slug", "n-1 PostCategory"],
            ["ProductCategory", "Id, Name, Description, Slug, ParentId, ImageUrl", "1-n Product; tự liên kết cha-con"],
            ["Product", "Id, Name, Description, Price, StockQuantity, TotalSold, ImageUrl, ProductCategoryId, Slug", "n-1 ProductCategory; 1-n OrderDetail"],
            ["Customer", "Id, FullName, Email, Phone, Password, TokenVersion, FirebaseUid, ResetPasswordToken", "1-n Order, CustomerAddress, Review"],
            ["Order", "Id, CustomerId, OrderDate, Status, PaymentMethod, PaymentStatus, ReceiverName, ShippingFee, DiscountAmount, TotalAmount", "n-1 Customer; 1-n OrderDetail"],
            ["OrderDetail", "Id, OrderId, ProductId, Quantity, UnitPrice", "n-1 Order; n-1 Product"],
            ["User", "Id, Username, PasswordHash, FullName, Role", "Tài khoản quản trị nội bộ"],
        ],
    )
    para(doc, "Cơ sở dữ liệu chạy trên SQL Server qua ApplicationDbContext. Kết quả dữ liệu thật tại thời điểm kiểm tra: 64 sản phẩm, 24 danh mục sản phẩm, 19 danh mục bài viết, 21 bài viết, 15 khách hàng, 49 đơn hàng, 81 dòng chi tiết đơn, 5 user admin và 154 cửa hàng.")
    heading(doc, "2.4. Danh mục Web API tiêu biểu", 2)
    add_table(
        doc,
        ["Nhóm", "Method", "Endpoint", "Mục đích"],
        [
            ["Products", "GET", "/api/products", "Danh sách sản phẩm có phân trang, lọc danh mục, tìm kiếm và sắp xếp."],
            ["Products", "GET", "/api/products/search", "Tìm kiếm theo từ khóa và khoảng giá."],
            ["Products", "GET", "/api/products/{id}", "Lấy chi tiết sản phẩm, ảnh và tùy chọn."],
            ["ProductCategories", "GET", "/api/product-categories/tree", "Lấy cây danh mục sản phẩm cho menu frontend."],
            ["Posts", "GET", "/api/posts", "Danh sách bài viết có phân trang/lọc danh mục."],
            ["Posts", "GET", "/api/posts/by-slug/{slug}", "Chi tiết bài viết theo slug SEO."],
            ["Customers", "POST", "/api/customers/register", "Đăng ký khách hàng, kiểm tra email trùng và hash mật khẩu."],
            ["Customers", "POST", "/api/customers/login", "Đăng nhập khách hàng, trả JWT token."],
            ["Orders", "POST", "/api/orders", "Tạo đơn hàng từ checkout frontend."],
            ["Orders", "GET", "/api/orders", "Lịch sử đơn hàng khách hàng có phân trang."],
            ["Stores", "GET", "/api/stores", "Danh sách cửa hàng phục vụ Store Locator."],
            ["Post", "POST", "/api/upload/ckeditor", "Upload ảnh từ CKEditor cho bài viết."],
        ],
    )

    # Chapter 3
    heading(doc, "CHƯƠNG 3: TRIỂN KHAI BACKEND ASP.NET CORE", 1)
    heading(doc, "3.1. Cấu trúc solution", 2)
    para(doc, "Solution NGUYENTRUCTRUONGCMS_SOLUTION.sln được tách thành các project CMS.Data, CMS.Backend, CMS.Frontend và các project kiểm thử. CMS.Data chứa entity, ApplicationDbContext và migration. CMS.Backend chứa controller MVC admin, controller API, service nghiệp vụ, cấu hình authentication/authorization, CORS và Swagger. CMS.Frontend là ứng dụng ReactJS chạy bằng Vite.")
    heading(doc, "3.2. Cấu hình database, migration và appsettings", 2)
    para(doc, "ApplicationDbContext khai báo đầy đủ DbSet cho các bảng nội dung, sản phẩm, khách hàng, đơn hàng, voucher, cửa hàng, đánh giá, banner và tùy chọn sản phẩm. Chuỗi kết nối SQL Server được cấu hình qua appsettings hoặc biến môi trường. Các khóa nhạy cảm như JWT key, SMTP/password gateway nên đặt ở user-secrets/biến môi trường khi triển khai thay vì ghi trực tiếp trong báo cáo.")
    heading(doc, "3.3. Middleware, CORS, Swagger và định tuyến", 2)
    para(doc, "Program.cs cấu hình CORS với policy AllowReactApp cho ReactJS localhost, đăng ký MapControllers để phục vụ Web API và giữ default MVC route cho trang quản trị. Swagger được bật để liệt kê nhóm API, schema DTO và hỗ trợ Bearer token cho các API cần xác thực.")
    add_image(doc, SWAGGER_REPORT, "Hình 3.1. Swagger Web API", width=6.3)
    add_image(doc, "02-api-products-json.png", "Hình 3.2. Kết quả JSON của API sản phẩm", width=5.8)
    heading(doc, "3.4. Xác thực và phân quyền quản trị", 2)
    para(doc, "Admin dùng Cookie Authentication. AccountController xác thực user bằng PasswordHasher, tạo claims gồm UserId, FullName và Role, sau đó chuyển vào Dashboard. UserController được bảo vệ bằng Authorize(Roles = \"Admin\") để chỉ tài khoản Admin được quản lý người dùng. Trang layout admin hiển thị tên người dùng và vai trò hiện tại.")
    add_image(doc, "03-admin-login.png", "Hình 3.3. Màn hình đăng nhập Admin", width=4.3)
    add_image(doc, "04-admin-dashboard.png", "Hình 3.4. Dashboard quản trị", width=4.3)
    heading(doc, "3.5. CRUD quản trị", 2)
    para(doc, "Các trang quản trị được xây dựng theo MVC/Razor, có sidebar chung, bộ lọc/tìm kiếm và phân trang. Nhóm chức năng trọng tâm gồm CRUD Category/Post/User, CategoryProduct/Product/Customer/Order/OrderDetail, quản lý cửa hàng, voucher, đánh giá và trạng thái đơn hàng.")
    add_image(doc, "05-admin-products-crud-content.png", "Hình 3.5. CRUD sản phẩm", width=4.3)
    add_image(doc, "06-admin-orders-crud-content.png", "Hình 3.6. CRUD đơn hàng", width=4.3)
    heading(doc, "3.6. CKEditor, upload ảnh và nội dung HTML", 2)
    para(doc, "Trang tạo/sửa bài viết tích hợp CKEditor để nhập nội dung giàu định dạng. API upload CKEditor nhận file ảnh, lưu vào wwwroot và trả đường dẫn cho editor. Khi frontend hiển thị chi tiết bài viết, React dùng dangerouslySetInnerHTML cho nội dung đã được quản trị biên tập.")
    add_image(doc, "07-admin-ckeditor-upload-content.png", "Hình 3.7. CKEditor và upload ảnh bài viết", width=4.3)
    heading(doc, "3.7. API đặt hàng, tồn kho, email và bảo mật mật khẩu", 2)
    para(doc, "OrderApiService nhận PlaceOrderDto từ checkout, khóa sản phẩm để kiểm tra tồn kho, tính subtotal, voucher, phí giao hàng, tạo Order/OrderDetail và trừ tồn kho sau khi tạo đơn thành công. EmailService gửi email chào mừng, xác nhận đơn hàng và reset password. Mật khẩu khách hàng dùng Microsoft.AspNetCore.Identity.PasswordHasher, tức cơ chế hash có salt/iteration tương đương chuẩn ASP.NET Identity, không lưu plaintext.")

    # Chapter 4
    heading(doc, "CHƯƠNG 4: TRIỂN KHAI FRONTEND REACTJS", 1)
    heading(doc, "4.1. Cấu trúc ReactJS và cấu hình API", 2)
    para(doc, "Frontend chạy bằng Vite, chia route theo Home, Menu, Product Detail, About/Post, Checkout, Auth, Profile và Stores. API base URL được cấu hình qua biến môi trường VITE_API_URL để không hardcode domain backend. Axios client chịu trách nhiệm gắn JWT token, gọi API JSON và xử lý lỗi.")
    heading(doc, "4.2. Trang chủ và các component", 2)
    para(doc, "Trang chủ được chia thành Header, HeroBanner, CategoryMenu, BestSellers, NewestProducts/ProductGrid và footer. Banner và danh sách sản phẩm lấy từ API, không dựng cứng nội dung. CategoryMenu hiển thị danh mục cấp cao và điều hướng người dùng vào trang shop/menu.")
    add_image(doc, "08-frontend-home-clean.png", "Hình 4.1. Trang chủ ReactJS", width=4.3)
    heading(doc, "4.3. Menu, tìm kiếm và lọc giá", 2)
    para(doc, "Trang menu/shop hỗ trợ tìm kiếm theo keyword, lọc theo danh mục và lọc theo khoảng giá. Khi người dùng thay đổi keyword hoặc giá min/max, React gọi lại API search/filter để cập nhật lưới sản phẩm. Trường hợp không có kết quả cần hiển thị thông báo phù hợp thay vì để màn hình trắng.")
    add_image(doc, "09-frontend-search-menu-clean.png", "Hình 4.2. Tìm kiếm/lọc sản phẩm", width=4.3)
    heading(doc, "4.4. Chi tiết sản phẩm, giỏ hàng và checkout", 2)
    para(doc, "Trang chi tiết sản phẩm hiển thị ảnh, mô tả, giá, trạng thái bán chạy/mới nhất, tùy chọn size, đường, đá và nút thêm giỏ hàng. Cart context quản lý danh sách sản phẩm, số lượng, tổng tiền và badge ở Header. Checkout bắt buộc khách nhập thông tin liên hệ/giao hàng trước khi gửi POST tạo đơn xuống backend.")
    add_image(doc, "10-product-detail-clean.png", "Hình 4.3. Chi tiết sản phẩm", width=6.3)
    add_image(doc, "11-product-added-cart-clean.png", "Hình 4.4. Thêm giỏ hàng và badge số lượng", width=4.3)
    add_image(doc, "12-checkout-form-clean.png", "Hình 4.5. Checkout", width=4.3)
    heading(doc, "4.5. Bài viết, HTML content và Store Locator", 2)
    para(doc, "Frontend có trang bài viết lấy nội dung HTML từ API và render đúng định dạng từ CKEditor. Store Locator lấy danh sách cửa hàng từ API, hiển thị danh sách và bản đồ với marker, giúp người dùng tìm cửa hàng gần vị trí mong muốn.")
    add_image(doc, "13-store-locator-map-clean.png", "Hình 4.6. Store locator và bản đồ", width=4.3)
    add_image(doc, "14-post-html-content-clean.png", "Hình 4.7. Chi tiết bài viết HTML", width=4.3)

    # Chapter 5
    heading(doc, "CHƯƠNG 5: KIỂM THỬ, MINH CHỨNG VÀ ĐÁNH GIÁ", 1)
    heading(doc, "5.1. Kết quả build và test", 2)
    add_table(
        doc,
        ["Hạng mục", "Lệnh/nguồn kiểm tra", "Kết quả"],
        [
            ["Backend build", "dotnet build NGUYENTRUCTRUONGCMS_SOLUTION.sln -c Release /p:UseAppHost=false", "Thành công, 0 warning, 0 error."],
            ["Backend test", "dotnet test NGUYENTRUCTRUONGCMS_SOLUTION.sln -c Release /p:UseAppHost=false --no-build", "29/29 test passed."],
            ["Frontend build", "npm run build", "Thành công; chỉ có cảnh báo chunk lớn của Vite."],
            ["Frontend lint", "npm run lint", "Thành công; 2 cảnh báo hook dependency, không có lỗi."],
            ["Swagger", "GET http://localhost:5188/swagger/index.html", "HTTP 200, hiển thị danh sách API và schemas."],
            ["Frontend runtime", "http://localhost:5173", "HTTP 200, console không có lỗi API/CORS trong các màn hình kiểm tra."],
        ],
    )
    heading(doc, "5.2. Kiểm thử luồng quản trị", 2)
    para(doc, "Đăng nhập admin hoạt động độc lập; nhập sai thông tin không làm crash ứng dụng. Các trang quản trị có sidebar, bộ lọc và bảng dữ liệu. Các action nhạy cảm được bảo vệ bởi Authorize, trong đó trang quản lý User yêu cầu role Admin.")
    heading(doc, "5.3. Kiểm thử luồng khách hàng", 2)
    para(doc, "Luồng frontend đã được chạy thực tế: mở trang chủ, tìm kiếm sản phẩm Latte, xem chi tiết sản phẩm, thêm vào giỏ hàng, thấy badge số lượng cập nhật, đi tới checkout, điền thông tin liên hệ và kiểm tra giao diện bản đồ cửa hàng. API JSON trả dữ liệu đúng định dạng success/data/items/totalItems.")
    heading(doc, "5.4. Đánh giá", 2)
    para(doc, "Dự án đáp ứng phần lớn tiêu chí chấm điểm: có cấu trúc solution 3 tầng, kết nối SQL Server bằng migration, quản trị CRUD, Web API GET/POST, Swagger, CORS, frontend ReactJS dùng API thật, giỏ hàng/checkout, upload CKEditor, hash mật khẩu, forgot password, email và ảnh minh chứng chạy thực tế. Một số điểm cần tiếp tục hoàn thiện trước khi nộp cuối là cập nhật lại README nếu thay đổi cổng chạy, bổ sung ảnh Postman nếu giảng viên yêu cầu riêng và xử lý triệt để hai cảnh báo dependency trong eslint.")

    # Chapter 6
    heading(doc, "CHƯƠNG 6: KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN", 1)
    heading(doc, "6.1. Kết luận", 2)
    para(doc, "Đồ án đã xây dựng được một hệ thống CMS và website bán hàng tương đối hoàn chỉnh theo kiến trúc ASP.NET Core + EF Core + SQL Server + ReactJS. Hệ thống có dữ liệu thật, API tài liệu hóa bằng Swagger, giao diện quản trị và giao diện khách hàng chạy được trên localhost, đồng thời có kiểm thử tự động để xác nhận các nghiệp vụ quan trọng.")
    heading(doc, "6.2. Hướng phát triển", 2)
    for item in [
        "Hoàn thiện thanh toán trực tuyến với callback thực tế và quy trình đối soát giao dịch.",
        "Bổ sung dashboard thống kê doanh thu, sản phẩm bán chạy và trạng thái tồn kho theo thời gian.",
        "Tăng coverage cho service đặt hàng, voucher, auth và frontend component tests.",
        "Triển khai lên môi trường cloud, tách cấu hình production bằng biến môi trường và bật HTTPS.",
        "Tối ưu UX mobile, tối ưu bundle frontend và xử lý cảnh báo lint còn lại.",
    ]:
        bullet(doc, item)

    # Appendix
    heading(doc, "PHỤ LỤC: ĐỐI CHIẾU TIÊU CHÍ CHẤM ĐIỂM", 1)
    checklist_rows = [
        ["1-3", "Solution 3 tầng, GitHub/README/.gitignore", "Đạt", "Repo có CMS.Data, CMS.Backend, CMS.Frontend; README hướng dẫn chạy; không phụ thuộc bin/obj."],
        ["4-9", "Báo cáo docx, 6 chương, ERD, API, Swagger", "Đạt", "Báo cáo này có 6 chương, ERD, bảng API, ảnh Swagger và JSON mẫu."],
        ["10-11", "8 entity, DbContext, migration, connection string", "Đạt", "ApplicationDbContext có đầy đủ entity; dữ liệu SQL Server chạy được."],
        ["12-15", "CRUD admin, phân trang, CKEditor", "Đạt", "Ảnh minh chứng CRUD sản phẩm/đơn hàng và CKEditor upload."],
        ["16-19", "Authorize, role Admin, layout, login/logout", "Đạt", "Cookie Auth, Authorize(Roles=\"Admin\"), layout hiển thị FullName/Role."],
        ["20-23", "Web API GET/POST, CORS, MapControllers + MVC", "Đạt", "Swagger liệt kê GET/POST; Program.cs cấu hình CORS/MapControllers/MVC."],
        ["24-32", "Frontend component, link, detail, cart, checkout, email, console", "Đạt", "Ảnh trang chủ/menu/detail/cart/checkout; Order API và EmailService hoạt động."],
        ["33-34", "Hash mật khẩu, register kiểm tra email trùng", "Đạt", "CustomerApiService dùng PasswordHasher và kiểm tra email trước khi tạo account."],
        ["35-36", "CKEditor upload, sản phẩm mới/bán chạy", "Đạt", "Có API upload CKEditor; Home gọi API sản phẩm mới/bán chạy."],
        ["37-45", "Category menu, lọc giá, search, badge cart, tồn kho, HTML content, env API URL", "Đạt", "Frontend dùng VITE_API_URL, search/filter/cart badge/detail HTML và kiểm soát tồn kho qua backend."],
        ["46", "Forgot password", "Đạt", "Có ForgotPassword/ResetPassword API, token hết hạn và email reset."],
        ["47-50", "Kỷ luật Git, teamwork, thuyết trình, tự học", "Cần bổ sung bằng lời khi nộp", "Nên thêm lịch sử commit/slide/thuyết trình nếu giảng viên yêu cầu riêng."],
    ]
    add_table(doc, ["Tiêu chí", "Nội dung", "Trạng thái", "Minh chứng"], checklist_rows)

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_document()
    print(OUTPUT)
